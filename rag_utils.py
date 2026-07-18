
"""
rag_utils.py (v1.0)
------------
All RAG (Retrieval-Augmented Generation) logic:
1. Extract text from an uploaded file (almost any common file type)
2. Split text into chunks
3. Convert chunks into embeddings
4. Store embeddings in a FAISS vector index
5. Search FAISS for the most relevant chunks for a question
6. Send question + retrieved context to a local Ollama LLM
 
SPEED NOTES (read this if answers feel slow):
- MODEL_NAME defaults to "llama3.2:1b" instead of the larger "llama3.2"
  (3B). The 1B model is noticeably faster on a normal laptop CPU and is
  still perfectly capable for document Q&A. If you prefer the previous
  model's slightly richer answers and don't mind waiting longer, change
  MODEL_NAME to "llama3.2" below (and run `ollama pull llama3.2` once).
- keep_alive="30m" is sent on every request so Ollama keeps the model
  loaded in memory between questions instead of reloading it each time
  (reloading is usually the single biggest source of slow answers).
- warm_up_model() is called once at server startup (see app.py) to
  preload the model *before* the first real question, so the user's
  first message isn't the slow one.
- num_predict caps how many tokens the model generates, which keeps
  answers focused and prevents unnecessarily long generation times.
- Chunk size/overlap and top_k were tuned down slightly so less text is
  sent to the model per question, which also speeds up generation.
"""
 
import os
import socket

# Force IPv4 DNS resolution to avoid EAI_NODATA / [Errno -5] errors on some hosts (like Render)
_orig_getaddrinfo = socket.getaddrinfo
def _patched_getaddrinfo(host, port, family=0, type=0, proto=0, flags=0):
    if family == 0 or family == socket.AF_UNSPEC:
        family = socket.AF_INET
    return _orig_getaddrinfo(host, port, family, type, proto, flags)
socket.getaddrinfo = _patched_getaddrinfo
import csv
import json
import re
import threading
import faiss
import numpy as np
import requests
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
# NOTE: sentence_transformers is intentionally NOT imported here at the top
# of the file. Just importing that library (before even loading a model)
# takes several seconds on its own, because it pulls in PyTorch. Since this
# file is imported by app.py before Flask starts, that import would still
# block startup even with the model loading itself moved to a background
# thread below. So the import is deferred to inside _load_embedding_model().
 
# ---------------------------------------------------------------------------
# 1. Load the embedding model in the BACKGROUND, not at import time.
#
# WHY: SentenceTransformer(...) used to be called directly here, at module
# level. That line runs the instant app.py does `from rag_utils import ...`
# -- which happens BEFORE Flask's app.run() is reached. Loading a PyTorch
# model (initializing torch, probing for a GPU, loading weights, building
# the tokenizer) takes several seconds on its own, so Flask couldn't start
# listening on port 5000 until that finished. This happened on EVERY
# launch, not just the first ever run -- unlike `pip install`, which is
# skipped after the first time. That's why the browser was always slow to
# connect, even after the "first run is slow" installs were long done.
#
# FIX: load it in a background thread (same pattern as warm_up_model()
# below for Ollama), so Flask can start accepting connections immediately.
# Anything that actually NEEDS the embedding model (uploading a file,
# asking a question) waits for it via _ensure_embedding_model_ready(),
# which is usually already finished by the time a real user gets that far.
# ---------------------------------------------------------------------------
embedding_model = None
reranker_model = None
_embedding_model_ready = threading.Event()

# bge-small-en-v1.5 is a much stronger retrieval model than the older
# all-MiniLM-L6-v2 at nearly the same size/speed -- this is what FINDS
# the right passages in a 100-200 page document. Same 384 dimensions.
EMBEDDING_MODEL_NAME = "BAAI/bge-small-en-v1.5"
EMBEDDING_DIM = 384
# bge models expect this prefix on the QUERY side (not on documents).
BGE_QUERY_PREFIX = "Represent this sentence for searching relevant passages: "

# Cross-encoder reranker: reads the question TOGETHER with each candidate
# passage and scores how well they truly match -- far more accurate than
# embedding similarity alone. Used as a second pass over the top results.
RERANKER_MODEL_NAME = "cross-encoder/ms-marco-MiniLM-L-6-v2"


class APIEmbeddingModel:
    """Lightweight embedding generator that calls Hugging Face's free Inference API."""
    def __init__(self, model_name="BAAI/bge-small-en-v1.5"):
        self.model_name = model_name
        self.api_url = f"https://api-inference.huggingface.co/pipeline/feature-extraction/{model_name}"
        self.token = os.getenv("HF_TOKEN", "")

    def encode(self, sentences, convert_to_numpy=True, normalize_embeddings=True, show_progress_bar=False, device=None):
        import requests
        import numpy as np
        
        headers = {}
        if self.token:
            headers["Authorization"] = f"Bearer {self.token}"
            
        if isinstance(sentences, str):
            sentences = [sentences]
            
        try:
            res = requests.post(
                self.api_url,
                headers=headers,
                json={"inputs": sentences, "options": {"wait_for_model": True}},
                timeout=60
            )
            res.raise_for_status()
            data = res.json()
            
            # Check for 3D list format [num_sentences, seq_len, dim] and mean pool if found
            if isinstance(data, list) and len(data) > 0:
                if isinstance(data[0], list) and len(data[0]) > 0 and isinstance(data[0][0], list):
                    pooled = []
                    for sent_emb in data:
                        seq_len = len(sent_emb)
                        dim = len(sent_emb[0])
                        mean_emb = [sum(tokens[d] for tokens in sent_emb) / seq_len for d in range(dim)]
                        pooled.append(mean_emb)
                    data = pooled
            
            embeddings = np.array(data, dtype="float32")
            
            if normalize_embeddings:
                norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
                norms = np.where(norms == 0, 1.0, norms)
                embeddings = embeddings / norms
                
            return embeddings
        except Exception as e:
            print(f"Hugging Face Inference API embeddings error: {e}")
            raise


def _load_embedding_model():
    global embedding_model, reranker_model
    if os.getenv("RENDER") == "true":
        print("Running on Render: using Hugging Face Inference API for embeddings.")
        embedding_model = APIEmbeddingModel(EMBEDDING_MODEL_NAME)
        reranker_model = None
        _embedding_model_ready.set()
        print("API Embedding model ready.")
        return

    print("Loading embedding + reranker models in the background...")
    # Deferred import -- see the NOTE near the top of this file for why.
    from sentence_transformers import SentenceTransformer, CrossEncoder
    # device="cpu" skips PyTorch's GPU-detection step, which can itself add
    # a noticeable delay on machines without a supported GPU.
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME, device="cpu")
    try:
        reranker_model = CrossEncoder(RERANKER_MODEL_NAME, device="cpu")
    except Exception as e:
        # Reranking is an accuracy bonus, not a requirement -- keep
        # working with embedding search alone if it can't load.
        print(f"Reranker not available ({e}); using embedding search only.")
        reranker_model = None
    _embedding_model_ready.set()
    print("Embedding model ready.")
 
 
_loading_started = False


def start_background_loading():
    """
    Kicks off the embedding-model load in a background thread. Called by
    app.py (instead of running automatically at import time) so that when
    Flask's auto-reloader is on, only the process that actually serves
    requests loads the model -- not the watcher process too.
    """
    global _loading_started
    if _loading_started:
        return
    _loading_started = True
    threading.Thread(target=_load_embedding_model, daemon=True).start()
 
 
def _ensure_embedding_model_ready(timeout=60):
    """
    Blocks (briefly) until the embedding model has finished loading in the
    background. In normal use this returns immediately, since the model
    usually finishes loading well before a user logs in and uploads a
    file. Only matters if someone uploads a file within the first couple
    of seconds of the server starting.
    """
    if not _embedding_model_ready.wait(timeout=timeout):
        raise RuntimeError(
            "The AI model is still starting up. Please wait a few seconds "
            "and try again."
        )
 
# How similar (cosine similarity, roughly -1 to 1) the best-matching chunk
# of the uploaded file must be to a question before we treat the question
# as "about the file". Below this, the question is answered from the
# model's own general knowledge instead. Tweak via .env if needed --
# lower catches more questions as file-related, higher is stricter.
#
# NOTE: 0.35 previously sat here and was too strict in practice -- broad
# questions like "what is in this file" or "how many rows are there"
# often don't share much vocabulary with any single chunk, so their score
# comes out low even though the file is clearly what's being asked about.
# 0.18 is a more realistic bar for "probably related"; the mentions_the_file()
# heuristic below catches broad/meta questions regardless of score.
#
# NOTE (bge upgrade): bge-small-en-v1.5 similarity scores run HIGHER
# than the old MiniLM model's -- unrelated text often scores ~0.5-0.6
# and related text 0.7+. The threshold is recalibrated accordingly.
# Wrongly routing a general question to the file is harmless anyway:
# the "no answer found" safety net retries from general knowledge.
RELEVANCE_THRESHOLD = float(os.environ.get("RELEVANCE_THRESHOLD", "0.6"))
 
# Broad or "meta" questions about the file itself -- summaries, counts,
# overviews -- should always be answered from the file if one is
# uploaded, even when their similarity score is low (see note above).
_FILE_REFERENCE_PATTERN = re.compile(
    r"(this file|the file|this document|the document|this data|the data|"
    r"this sheet|the sheet|this spreadsheet|the spreadsheet|this table|"
    r"uploaded file|my file|attached file|in (this|the) file|"
    r"in (this|the) document|summari|overview|how many|"
    r"what('?s| is) in|tell me about|"
    # Image/scan uploads: "what does the image say", "read the photo"...
    r"(this|the|my) (image|photo|picture|scan|screenshot|pic)|"
    r"what (does|do|is) (it|this|that|the \w+) say|"
    r"what('?s| is) written|read (it|this|that) (out|for me)|"
    # Marksheet/certificate-style questions
    r"marks|grade|score|total|result|percentage|amount)",
    re.IGNORECASE,
)
 
# If a file-grounded answer still sounds like it came up empty, that's a
# sign the retrieved chunks genuinely weren't useful -- a safety net for
# falling back to general knowledge instead of leaving a dead-end answer.
_NO_ANSWER_FOUND_PATTERN = re.compile(
    r"(i (don'?t|do not) have (any )?information|"
    r"i (don'?t|do not) (know|see) (any|the)|"
    r"i couldn'?t find|i could not find|"
    r"no (information|mention) (of|about)|"
    r"haven'?t (seen|uploaded)|"
    r"not (mentioned|found|available) in (the|this) (file|document|context))",
    re.IGNORECASE,
)
 
 
def mentions_the_file(question):
    """
    Heuristic: does this question sound like it's asking about the
    uploaded file itself (a summary/count/overview), as opposed to an
    unrelated general-knowledge question? Used to route broad questions
    to the file even when they score low on similarity search.
    """
    return bool(_FILE_REFERENCE_PATTERN.search(question))
 
 
def sounds_like_no_answer_found(answer):
    """
    Heuristic: does the model's own answer sound like it doesn't
    actually have the context, rather than that phrasing appearing as a
    normal factual answer? Used as a safety net after answering from the
    file, to retry from general knowledge if the file path came up empty.
    """
    return bool(_NO_ANSWER_FOUND_PATTERN.search(answer))
 
# ---------------------------------------------------------------------------
# LLM BACKEND SELECTION: Groq (cloud) or Ollama (local)
# ---------------------------------------------------------------------------
# The app can answer using either:
#   * Groq   -- a fast, free cloud API (used when GROQ_API_KEY is set). This
#               is what makes the app work when hosted online, where there's
#               no local Ollama running.
#   * Ollama -- a model running locally on this machine (the original mode,
#               used when no GROQ_API_KEY is present).
# Whichever is active is chosen automatically at startup based on the
# environment, so the same code runs locally AND when deployed to the web.
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "").strip()
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
# openai/gpt-oss-20b: a current Groq production model -- very fast (~1000
# tokens/sec), cheap, 131K context. Override via GROQ_MODEL in .env.
GROQ_MODEL = os.environ.get("GROQ_MODEL", "openai/gpt-oss-20b")

USE_GROQ = bool(GROQ_API_KEY)

# ---------------------------------------------------------------------------
# Ollama settings (local backend)
# ---------------------------------------------------------------------------
OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = os.environ.get("OLLAMA_MODEL", "llama3.2:1b")

# Reuse one HTTP connection instead of opening a new one per request.
_session = requests.Session()
 
 
def warm_up_model():
    """
    Sends a tiny throwaway prompt to Ollama at startup so the model gets
    loaded into memory before any real user question arrives. Runs in a
    background thread so it never blocks the server from starting.
    """
    # Groq is a cloud API with no cold-start on our side -- nothing to warm.
    if USE_GROQ:
        print(f"Using Groq cloud model '{GROQ_MODEL}' -- no local warm-up needed.")
        return

    def _warm():
        try:
            print(f"Warming up Ollama model '{MODEL_NAME}'...")
            _session.post(
                OLLAMA_URL,
                json={
                    "model": MODEL_NAME,
                    "prompt": "Hello",
                    "stream": False,
                    "keep_alive": "30m",
                    "options": {"num_predict": 1},
                },
                timeout=120,
            )
            print("Model warm-up complete. Answers should now be fast.")
        except Exception as e:
            print(f"Model warm-up skipped (Ollama may not be running yet): {e}")
 
    threading.Thread(target=_warm, daemon=True).start()
 
 
# ---------------------------------------------------------------------------
# STEP 1: TEXT EXTRACTION
# ---------------------------------------------------------------------------
def extract_text(file_path):
    """
    Reads almost any common file type and returns its content as plain text.
    Supported directly: PDF, DOCX, TXT, MD, CSV, JSON, XLSX, XLS.
    Anything else: tried as plain text; if that fails, a clear error is
    raised instead of a crash.
    """
    extension = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""
    text = ""
 
    if extension == "pdf":
        text = extract_pdf_text(file_path)
 
    elif extension == "docx":
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                text += ", ".join(cell.text for cell in row.cells) + "\n"
 
    elif extension in ("txt", "md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
 
    elif extension == "csv":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                text += ", ".join(row) + "\n"
 
    elif extension == "xlsx":
        text = extract_excel_text(file_path, engine="openpyxl")
 
    elif extension == "xls":
        text = extract_excel_text(file_path, engine="xlrd")
 
    elif extension == "json":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
        text = json.dumps(data, indent=2)
 
    elif extension == "ipynb":
        text = extract_ipynb_text(file_path)
 
    elif extension in ("jpg", "jpeg", "png", "bmp", "tiff", "webp"):
        text = extract_image_text(file_path)
 
    else:
        try:
            with open(file_path, "r", encoding="utf-8", errors="strict") as f:
                text = f.read()
        except Exception:
            raise ValueError(
                f"I can't read text from '.{extension}' files (this is "
                "usually an image or another non-text format). Please "
                "upload a PDF, Word, Excel, CSV, TXT, Markdown, or JSON "
                "file instead."
            )
 
    if not text.strip():
        raise ValueError("No readable text was found in this file.")
 
    return text
 
 
def extract_pdf_text(file_path):
    """
    Extracts text from a PDF using two libraries for reliability:
    PyPDF2 first, then pdfplumber as a fallback for awkward PDFs.
    """
    text = ""
 
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception:
        text = ""
 
    if len(text.strip()) < 20:
        try:
            import pdfplumber
            fallback_text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        fallback_text += page_text + "\n"
            if len(fallback_text.strip()) > len(text.strip()):
                text = fallback_text
        except Exception:
            pass
 
    if len(text.strip()) < 5:
        # No text layer -- this is a scanned/image PDF. Fall back to OCR:
        # render each page to an image and read the text out of it.
        text = ocr_pdf_pages(file_path)
        if len(text.strip()) < 5:
            raise ValueError(
                "OCR ran on this scanned PDF but couldn't find readable "
                "text. The scan may be too blurry or low-contrast -- try "
                "a clearer scan or a text-based PDF."
            )
 
    return text
 
 
def _preprocess_for_ocr(image):
    """
    Cleans an image up before OCR to noticeably improve accuracy on
    phone photos and scans: grayscale (removes color noise), auto-
    contrast (makes faded text darker), and upscaling small images
    (Tesseract reads small text badly).
    """
    from PIL import Image, ImageOps

    image = ImageOps.grayscale(image)
    image = ImageOps.autocontrast(image, cutoff=1)
    # Upscale if the image is small -- OCR accuracy drops sharply on
    # low-resolution photos (e.g. WhatsApp-compressed images).
    if image.width < 1500:
        scale = 1500 / image.width
        image = image.resize(
            (int(image.width * scale), int(image.height * scale)),
            Image.LANCZOS,
        )
    return image


def _configure_tesseract(pytesseract):
    """
    Makes sure pytesseract can find the Tesseract program even when it
    isn't on PATH (common right after installing on Windows -- the
    installer puts it in Program Files but doesn't always update PATH
    for already-open windows).
    """
    import shutil
    if shutil.which("tesseract"):
        return
    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"),
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
        "/opt/homebrew/bin/tesseract",
    ]
    for path in candidates:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            return
 
 
def ocr_pdf_pages(file_path):
    """
    OCR fallback for scanned PDFs (no text layer): renders each page to
    an image with pypdfium2 (already installed -- it ships with
    pdfplumber) and reads the text with Tesseract OCR.
    """
    try:
        import pytesseract
        from PIL import Image  # noqa: F401 -- needed by bitmap.to_pil()
        import pypdfium2 as pdfium
    except ImportError:
        raise ValueError(
            "This PDF is a scanned image, so reading it needs OCR -- but "
            "the OCR packages aren't installed yet. Re-run run.bat (or "
            "`pip install -r requirements.txt`) to install them, and make "
            "sure the Tesseract program is installed (see the README's "
            "OCR section)."
        )
 
    _configure_tesseract(pytesseract)
 
    try:
        pdf = pdfium.PdfDocument(file_path)
        pages_text = []
        # scale 300/72 renders at ~300 DPI, a good resolution for OCR.
        for page in pdf:
            image = page.render(scale=300 / 72).to_pil()
            image = _preprocess_for_ocr(image)
            pages_text.append(pytesseract.image_to_string(image))
        return "\n".join(pages_text)
    except pytesseract.TesseractNotFoundError:
        raise ValueError(
            "This PDF is a scanned image, so reading it needs the "
            "Tesseract OCR program, which wasn't found on this computer. "
            "Install it from https://github.com/UB-Mannheim/tesseract/wiki "
            "(Windows) and try again -- see the README's OCR section."
        )
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"OCR failed on this scanned PDF. (Details: {e})")
 
 
def extract_image_text(file_path):
    """
    Extracts text from an image (JPG, PNG, etc.) using OCR, if OCR support
    is installed. OCR requires an extra SYSTEM-level program (Tesseract),
    not just a Python package, so it's kept fully optional:
    - If pytesseract + Tesseract are installed, images work like any
      other file (see README for the one-time setup step).
    - If not installed, uploading an image gives a clear, friendly
      message instead of a crash or a confusing traceback.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ValueError(
            "Reading text from images requires an optional OCR feature "
            "that isn't set up yet. See the README section 'Enabling "
            "image (OCR) support' for a short one-time setup step. "
            "In the meantime, please upload a PDF, Word, Excel, CSV, "
            "TXT, or JSON file instead."
        )
 
    _configure_tesseract(pytesseract)

    try:
        image = Image.open(file_path)
        image = _preprocess_for_ocr(image)
        text = pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError:
        raise ValueError(
            "OCR is installed as a Python package, but the underlying "
            "Tesseract program wasn't found on this computer. See the "
            "README section 'Enabling image (OCR) support' to install it."
        )
    except Exception as e:
        raise ValueError(f"Could not read text from this image. (Details: {e})")
 
    if not text.strip():
        raise ValueError(
            "No readable text was found in this image. OCR works best on "
            "clear, high-contrast text -- a photo of a screen or a blurry "
            "scan may not extract cleanly."
        )
 
    return text
 
 
def extract_excel_text(file_path, engine):
    """
    Reads every sheet of an Excel file into readable text. Tries the
    engine matching the file's extension first; if that fails (a common
    cause: a file saved with the wrong extension, e.g. an old .xls file
    renamed to .xlsx), automatically retries with the other engine before
    giving up -- the same "try a second way before failing" approach used
    for PDFs.
 
    A very common real-world case this also handles: many "Excel" exports
    from banks, invoicing tools, and web dashboards are not real Excel
    binaries at all -- they're an HTML table saved with an .xls/.xlsx
    extension (this is what causes errors like "Expected BOF record,
    found b'<html>'"). If both real-Excel engines fail, we fall back to
    reading the file as an HTML table before giving up entirely.
    """
    other_engine = "xlrd" if engine == "openpyxl" else "openpyxl"
 
    try:
        sheets = pd.read_excel(file_path, sheet_name=None, engine=engine)
    except Exception as first_error:
        try:
            sheets = pd.read_excel(file_path, sheet_name=None, engine=other_engine)
        except Exception:
            # Last resort: this "Excel" file might actually be HTML.
            try:
                tables = pd.read_html(file_path)
                if not tables:
                    raise ValueError("no tables found")
                sheets = {f"Table {i + 1}": df for i, df in enumerate(tables)}
            except Exception:
                raise ValueError(
                    f"Could not read this Excel file. It may be corrupted, "
                    f"password-protected, or saved in an unusual format. "
                    f"(Details: {first_error})"
                )
 
    text = ""
    for sheet_name, df in sheets.items():
        text += _dataframe_to_text(df, sheet_name)
    return text


def _dataframe_to_text(df, sheet_name):
    """
    Turns a spreadsheet into text the AI can actually reason about.

    Real-world exports (reports from pharma/CRM/banking tools) often have
    several junk rows (logos, titles, date ranges) BEFORE the real column
    header row, so pandas reads meaningless numeric column names. This:
    1. Detects the real header row within the first ~10 rows.
    2. Writes each data row as labeled "Column: value" pairs -- so a
       question like "how many visit dates in one row" or "what is
       ALKA MAKWANA's specialty" maps directly onto the text.
    """
    df = df.copy()

    def _is_header_like(row_vals):
        vals = [str(v).strip() for v in row_vals]
        filled = [v for v in vals if v and v.lower() != "nan"]
        if len(filled) < max(3, int(len(vals) * 0.5)):
            return False
        # Header cells are text labels, not numbers, and mostly unique.
        numericish = sum(1 for v in filled if v.replace(".", "", 1).replace("-", "").isdigit())
        return numericish <= 1 and len(set(filled)) >= len(filled) - 1

    # If pandas already found meaningful string headers, keep them.
    cols_meaningless = all(
        str(c).strip().isdigit() or str(c).lower().startswith("unnamed")
        for c in df.columns
    )
    if cols_meaningless:
        for i in range(min(10, len(df))):
            if _is_header_like(df.iloc[i].tolist()):
                df.columns = [str(v).strip() for v in df.iloc[i]]
                df = df.iloc[i + 1:].reset_index(drop=True)
                break

    header = [str(c).strip() for c in df.columns]
    lines = [
        f"Sheet: {sheet_name}",
        f"This sheet has {len(df)} data rows. Columns: {', '.join(h for h in header if h and h.lower() != 'nan')}.",
        "",
    ]
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        pairs = []
        for col, val in zip(header, row.tolist()):
            sval = str(val).strip()
            if not sval or sval.lower() == "nan" or not col or col.lower() == "nan":
                continue
            pairs.append(f"{col}: {sval[:120]}")
        if pairs:
            lines.append(f"Row {i}: " + " | ".join(pairs))
    return "\n".join(lines) + "\n\n"
 
 
def extract_ipynb_text(file_path):
    """
    Reads a Jupyter notebook (.ipynb) -- these are plain JSON under the
    hood -- and pulls out markdown text, code cells, and any text-based
    outputs (print statements, printed dataframes, etc.) so the notebook
    can be asked about like any other document.
    """
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            notebook = json.load(f)
    except json.JSONDecodeError:
        raise ValueError(
            "This .ipynb file doesn't look like a valid Jupyter notebook "
            "(it couldn't be parsed as JSON)."
        )
 
    def _join(source):
        if isinstance(source, list):
            return "".join(source)
        return source or ""
 
    parts = []
    for cell in notebook.get("cells", []):
        cell_type = cell.get("cell_type")
        content = _join(cell.get("source", "")).strip()
        if not content:
            continue
 
        if cell_type == "markdown":
            parts.append(content)
        elif cell_type == "code":
            parts.append("Code:\n" + content)
            for output in cell.get("outputs", []):
                out_text = ""
                if "text" in output:
                    out_text = _join(output["text"])
                elif "data" in output and "text/plain" in output["data"]:
                    out_text = _join(output["data"]["text/plain"])
                out_text = out_text.strip()
                if out_text:
                    parts.append("Output:\n" + out_text)
 
    return "\n\n".join(parts)
 
 
# ---------------------------------------------------------------------------
# STEP 2: CHUNKING
# ---------------------------------------------------------------------------
def chunk_text(text, chunk_size=1000, overlap_sentences=2):
    """
    Splits long text into overlapping chunks WITHOUT cutting through the
    middle of sentences. The old version sliced every 600 characters
    blindly, which often split a sentence (or a number!) in half -- on a
    100-200 page document that noticeably hurt search quality, because a
    chunk that starts mid-sentence embeds poorly and reads poorly.

    Approach: split the text into sentences, then pack whole sentences
    into ~chunk_size-character chunks. Each new chunk repeats the last
    couple of sentences of the previous one, so facts that straddle a
    boundary still appear intact in at least one chunk.
    """
    # Split into sentences (handles ". ", "! ", "? ", and newlines).
    # OCR text can lack punctuation, so long newline-separated lines are
    # treated as sentence breaks too.
    sentences = re.split(r"(?<=[.!?])\s+|\n{2,}", text)
    sentences = [s.strip() for s in sentences if s and s.strip()]

    chunks = []
    current = []
    current_len = 0

    for sentence in sentences:
        # A single monster "sentence" (e.g. an OCR'd table) gets hard-split.
        while len(sentence) > chunk_size:
            head, sentence = sentence[:chunk_size], sentence[chunk_size:]
            if current:
                chunks.append(" ".join(current))
                current, current_len = [], 0
            chunks.append(head.strip())

        if current_len + len(sentence) > chunk_size and current:
            chunks.append(" ".join(current))
            # Overlap: carry the last few sentences into the next chunk.
            current = current[-overlap_sentences:] if overlap_sentences else []
            current_len = sum(len(s) + 1 for s in current)
        current.append(sentence)
        current_len += len(sentence) + 1

    if current:
        chunks.append(" ".join(current))

    return [c for c in (c.strip() for c in chunks) if c]


# ---------------------------------------------------------------------------
# STEP 3 & 4: EMBEDDINGS + FAISS VECTOR STORAGE (with disk persistence)
# ---------------------------------------------------------------------------
class VectorStore:
    """
    A FAISS index + the text chunks it represents, for one user's one
    active document. Can save/load itself to disk so a server restart
    doesn't force the user to re-upload and re-index their file.
    """

    def __init__(self):
        self.index = None
        self.chunks = []

    def build(self, chunks):
        """Convert chunks to embeddings and build a fresh FAISS index."""
        _ensure_embedding_model_ready()
        self.chunks = chunks
        # normalize_embeddings=True turns this into unit vectors, so an
        # inner-product index below gives us cosine similarity scores in
        # a predictable, easy-to-threshold [-1, 1] range (used to decide
        # whether a question is actually about the uploaded file).
        embeddings = embedding_model.encode(chunks, convert_to_numpy=True, normalize_embeddings=True)
        embeddings = np.array(embeddings, dtype="float32")

        self.index = faiss.IndexFlatIP(EMBEDDING_DIM)
        self.index.add(embeddings)

    def search(self, question, top_k=3):
        """
        Finds the most relevant chunks for the given question, in two
        stages (this is what makes answers accurate on BIG documents):

        1. RETRIEVE: embedding search pulls a wide pool of candidate
           chunks (fast, approximate).
        2. RERANK: the cross-encoder re-reads the question against each
           candidate and re-orders them by true relevance (slower but
           precise) -- the best top_k survive.

        Returns (chunks, scores) where scores are the embedding cosine
        similarities of the returned chunks (used for the relevance
        threshold check in app.py).
        """
        if self.index is None or self.index.ntotal == 0:
            return [], []

        _ensure_embedding_model_ready()
        question_vector = embedding_model.encode(
            [BGE_QUERY_PREFIX + question],
            convert_to_numpy=True, normalize_embeddings=True,
        )
        question_vector = np.array(question_vector, dtype="float32")

        # Stage 1: pull a wide candidate pool (more than we need).
        pool_size = max(top_k * 4, 20) if reranker_model is not None else top_k
        pool_size = min(pool_size, self.index.ntotal)
        scores, indices = self.index.search(question_vector, pool_size)

        candidates = []
        for score, idx in zip(scores[0], indices[0]):
            if 0 <= idx < len(self.chunks):
                candidates.append((self.chunks[idx], float(score)))

        if not candidates:
            return [], []

        # Stage 2: rerank the pool with the cross-encoder.
        if reranker_model is not None and len(candidates) > top_k:
            try:
                pairs = [(question, chunk) for chunk, _ in candidates]
                rerank_scores = reranker_model.predict(pairs)
                order = sorted(
                    range(len(candidates)),
                    key=lambda i: float(rerank_scores[i]),
                    reverse=True,
                )
                candidates = [candidates[i] for i in order[:top_k]]
            except Exception:
                candidates = candidates[:top_k]
        else:
            candidates = candidates[:top_k]

        results = [c for c, _ in candidates]
        result_scores = [s for _, s in candidates]
        return results, result_scores

    def save(self, folder):
        """Persist the index + chunks to disk so they survive a restart."""
        if self.index is None:
            return
        os.makedirs(folder, exist_ok=True)
        faiss.write_index(self.index, os.path.join(folder, "index.faiss"))
        with open(os.path.join(folder, "chunks.json"), "w", encoding="utf-8") as f:
            json.dump(self.chunks, f)
        # Record which embedding model built this index -- an index built
        # by a different model is meaningless to search with.
        with open(os.path.join(folder, "index_meta.json"), "w", encoding="utf-8") as f:
            json.dump({"embedding_model": EMBEDDING_MODEL_NAME}, f)

    def load(self, folder):
        """
        Loads a previously saved index. Returns True if a saved index was
        found, loaded, AND was built with the current embedding model --
        otherwise False (the app will ask for a re-upload).
        """
        index_path = os.path.join(folder, "index.faiss")
        chunks_path = os.path.join(folder, "chunks.json")
        meta_path = os.path.join(folder, "index_meta.json")
        if not (os.path.exists(index_path) and os.path.exists(chunks_path)):
            return False

        # Reject indexes built by an older/different embedding model.
        if os.path.exists(meta_path):
            try:
                with open(meta_path, "r", encoding="utf-8") as f:
                    meta = json.load(f)
                if meta.get("embedding_model") != EMBEDDING_MODEL_NAME:
                    return False
            except Exception:
                return False
        else:
            # No meta file = index from before this upgrade.
            return False

        self.index = faiss.read_index(index_path)
        with open(chunks_path, "r", encoding="utf-8") as f:
            self.chunks = json.load(f)
        return True


# ---------------------------------------------------------------------------
# STEP 5: CALL LOCAL OLLAMA LLM
# ---------------------------------------------------------------------------
# Cap how much document context is sent to the model per question.
# Prompt processing on a CPU is a big part of answer latency, so a
# smaller, focused context is dramatically faster (and the small model
# handles short contexts better anyway).
# Balance point: enough context for accurate answers, small enough that
# a CPU can read the prompt quickly. (Reranking means the chunks that
# make it in are the RIGHT ones, so we don't need to send as many.)
MAX_CONTEXT_CHARS = 5000


def _build_prompt(question, context_chunks=None):
    """Builds the LLM prompt (file-grounded or general-knowledge)."""
    if context_chunks:
        # Trim context to the cap, keeping whole chunks where possible.
        trimmed, total = [], 0
        for c in context_chunks:
            if total + len(c) > MAX_CONTEXT_CHARS and trimmed:
                break
            trimmed.append(c)
            total += len(c)
        context_text = "\n\n".join(f"- {c}" for c in trimmed)
        prompt = f"""You are a helpful assistant answering questions about a document
the user has uploaded. Answer naturally and clearly, the way a knowledgeable
person would explain it out loud — in complete sentences, not just copied
fragments or a list of keywords.

Guidelines:
- Base your answer only on the context below.
- If the question is broad (e.g. "what is this document about"), summarize
  the key points from the context in your own words.
- If the exact answer isn't stated but something closely related is in the
  context, explain what IS there instead of simply saying it's missing.
- Only say you couldn't find the answer if the context is genuinely
  unrelated to the question.
- The context may come from OCR of a scanned image or photo, so it can
  contain garbled characters or odd spacing -- ignore the noise and
  focus on the meaningful text.
- OCR cannot read currency signs: a rupee sign usually comes out as
  #, @, or Z (so "#249" almost certainly means 249 rupees).
- Payment/app screenshots mix ADVERTISEMENTS with the real transaction.
  Numbers inside offer sentences ("get up to X", "X cashback", reward
  gimmicks like "1=1paisa") are NEVER the paid amount. The real amount
  is the standalone number near the recipient's name, "Payment
  successful", or the date.
  WORKED EXAMPLE -- given OCR text:
    "Get up to @1,000 on every payment @1=1paisa Payment successful
     to RAKESH KUMAR #350 Paid via UPI 2 Jan 2026"
  the amount actually paid to Rakesh is 350 rupees. The 1,000 and the
  1=1paisa are advertising and must be ignored.
- Keep the answer clear and reasonably concise, but complete.

Context from the document:
{context_text}

User's question: {question}

Answer:"""
    else:
        prompt = f"""You are a helpful, knowledgeable assistant.
Answer the user's question clearly and naturally, in complete sentences.
Do not comment on the question itself -- just answer it. If you don't
know something, say so honestly instead of guessing.

User's question: {question}

Answer:"""
    return prompt


_OLLAMA_OPTIONS = {
    # Caps answer length -- 400 tokens is enough for a thorough,
    # human-sounding answer without rambling.
    "num_predict": 400,
    # Context window sized for MAX_CONTEXT_CHARS of document text plus
    # the prompt and answer.
    "num_ctx": 4096,
}

# "Thinking" models (qwen3, deepseek-r1, ...) spend a long time reasoning
# to themselves before answering -- great for math, terrible for waiting.
# Disable it so document Q&A answers start immediately. The flag is only
# sent to models that support it (others would reject the request).
_THINKING_MODELS = ("qwen3", "deepseek-r1", "gpt-oss", "magistral")
_EXTRA_PAYLOAD = (
    {"think": False}
    if any(name in MODEL_NAME.lower() for name in _THINKING_MODELS)
    else {}
)


def ask_ollama(question, context_chunks=None):
    """
    Sends a question to the locally running Ollama model and returns the
    complete answer in one go (non-streaming).
    """
    prompt = _build_prompt(question, context_chunks)

    try:
        response = _session.post(
            OLLAMA_URL,
            json={
                "model": MODEL_NAME,
                "prompt": prompt,
                "stream": False,
                # Keeps the model loaded in memory between requests instead
                # of reloading it each time -- the single biggest speed win.
                "keep_alive": "30m",
                "options": _OLLAMA_OPTIONS,
                **_EXTRA_PAYLOAD,
            },
            timeout=120,
        )
    except requests.exceptions.ConnectionError:
        raise RuntimeError(
            "Could not connect to Ollama. Make sure the Ollama app is "
            "running in the background (it should start automatically "
            "after installation, or run 'ollama serve' manually)."
        )

    if response.status_code != 200:
        raise RuntimeError(f"Ollama error: {response.text}")

    return response.json().get("response", "").strip()


# ---------------------------------------------------------------------------
# STEP 5 (cloud): CALL GROQ LLM  -- used when GROQ_API_KEY is set
# ---------------------------------------------------------------------------
def ask_groq(question, context_chunks=None):
    """
    Sends a question to Groq's cloud API and returns the complete answer
    in one go (non-streaming). Uses the OpenAI-compatible chat endpoint.
    """
    prompt = _build_prompt(question, context_chunks)
    try:
        response = _session.post(
            GROQ_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": GROQ_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 800,
                "stream": False,
            },
            timeout=120,
        )
    except requests.exceptions.ConnectionError:
        raise RuntimeError(
            "Could not reach Groq. Check the server's internet connection "
            "and that GROQ_API_KEY is set correctly."
        )

    if response.status_code != 200:
        raise RuntimeError(f"Groq error ({response.status_code}): {response.text}")

    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


def ask_groq_stream(question, context_chunks=None):
    """
    Same as ask_groq, but yields the answer piece by piece as Groq
    generates it (Server-Sent Events), so the UI shows text immediately.
    """
    prompt = _build_prompt(question, context_chunks)
    try:
        with _session.post(
            GROQ_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": GROQ_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 800,
                "stream": True,
            },
            stream=True,
            timeout=300,
        ) as response:
            if response.status_code != 200:
                raise RuntimeError(
                    f"Groq error ({response.status_code}): {response.text}"
                )
            for line in response.iter_lines():
                if not line:
                    continue
                line = line.decode("utf-8")
                if not line.startswith("data: "):
                    continue
                payload = line[len("data: "):].strip()
                if payload == "[DONE]":
                    break
                try:
                    data = json.loads(payload)
                    piece = data["choices"][0]["delta"].get("content", "")
                    if piece:
                        yield piece
                except (json.JSONDecodeError, KeyError, IndexError):
                    continue
    except requests.exceptions.ConnectionError:
        raise RuntimeError(
            "Could not reach Groq. Check the server's internet connection "
            "and that GROQ_API_KEY is set correctly."
        )


def ask_ollama_stream(question, context_chunks=None):
    """
    Same as ask_ollama, but yields the answer piece by piece as the
    model generates it, so the UI can show text immediately instead of
    waiting for the whole answer -- this is what makes responses FEEL
    fast even on a slower computer.
    """
    prompt = _build_prompt(question, context_chunks)

    try:
        with _session.post(
            OLLAMA_URL,
            json={
                "model": MODEL_NAME,
                "prompt": prompt,
                "stream": True,
                "keep_alive": "30m",
                "options": _OLLAMA_OPTIONS,
                **_EXTRA_PAYLOAD,
            },
            stream=True,
            timeout=300,
        ) as response:
            if response.status_code != 200:
                raise RuntimeError(f"Ollama error: {response.text}")
            for line in response.iter_lines():
                if not line:
                    continue
                data = json.loads(line.decode("utf-8"))
                piece = data.get("response", "")
                if piece:
                    yield piece
                if data.get("done"):
                    break
    except requests.exceptions.ConnectionError:
        raise RuntimeError(
            "Could not connect to Ollama. Make sure the Ollama app is "
            "running in the background (it should start automatically "
            "after installation, or run 'ollama serve' manually)."
        )


# ---------------------------------------------------------------------------
# BACKEND DISPATCHERS -- app.py calls these; they route to Groq or Ollama.
# ---------------------------------------------------------------------------
def ask_llm(question, context_chunks=None):
    """Answer using whichever backend is active (Groq if configured, else Ollama)."""
    if USE_GROQ:
        return ask_groq(question, context_chunks)
    return ask_ollama(question, context_chunks)


def ask_llm_stream(question, context_chunks=None):
    """Streaming answer via whichever backend is active."""
    if USE_GROQ:
        yield from ask_groq_stream(question, context_chunks)
    else:
        yield from ask_ollama_stream(question, context_chunks)
